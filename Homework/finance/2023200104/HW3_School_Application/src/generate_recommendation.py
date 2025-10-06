import json, os
from docx import Document


def load_json(path):
    try:
        with open(path, 'r', encoding='utf-8-sig') as f:
            return json.load(f)
    except Exception:
        return {}

base = r"C:\Users\32854\Desktop\quant课\HW_School_Application"
out_dir = os.path.join(base, "output", "letters")
os.makedirs(out_dir, exist_ok=True)
info = load_json(os.path.join(base, "personal_info.json"))

name_cn = "吴昭毅"
name_en = "Wuzhaoyi (Jeffri) Wu"
gpa = "3.99"
major = "Financial Engineering"
email = "3285448436@qq.com"
phone = "16688308888"
background = "quantitative trading"

name_cn = info.get("name_cn", name_cn)
name_en = info.get("name_en", name_en)
gpa = str(info.get("gpa", gpa))
major = info.get("major", major)
email = info.get("email", email)
phone = info.get("phone", phone)
background = info.get("background_summary", background)

# Recommender details
recommender_name_en = info.get("recommender_name_en", "Lei Ge")
recommender_name_cn = info.get("recommender_name_cn", "葛雷")
recommender_title = info.get("recommender_title", "EnnGroup Assistant Professor in Quantitative Economics")
recommender_affiliation = info.get("recommender_affiliation", "Renmin University of China")
recommender_email = info.get("recommender_email", "")
recommender_phone = info.get("recommender_phone", "")


doc = Document()

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

add_para("To the Admissions Committee:", bold=False)
add_para("")
body = f"""As {recommender_title} at {recommender_affiliation}, I am delighted to recommend {name_en} for admission to your graduate program in Financial Engineering or Quantitative Finance. I have observed {name_en.split()[0]}'s academic performance and professional growth closely and consider {name_en.split()[0]} an exceptional candidate with strong...

{name_en.split()[0]} has maintained an outstanding academic record, reflected in a GPA of {gpa} in {major}. Coursework and independent study demonstrate mastery in probability and statistics, stochastic processes, optimization, econometrics, and computational methods. {name_en.split()[0]} consistent...

Technically, {name_en.split()[0]} is a highly capable quantitative developer. Fluent in Python and the scientific computing stack (including NumPy, pandas, and common machine learning libraries), {name_en.split()[0]} builds clean, modular, and reproducible codebases under version control. In project...

Beyond technical competence, {name_en.split()[0]} communicates complex ideas with clarity. {name_en.split()[0]} writes with precision, presents results logically, and is receptive to critique iterating quickly while maintaining methodological integrity. In team settings, {name_en.split()[0]} collabor...

Given {name_en.split()[0]}'s academic excellence, quantitative rigor, and professional maturity, I recommend {name_en} to your program without reservation. I am confident that {name_en.split()[0]} will contribute meaningfully to your academic community and excel in advanced study and research.
"""
for para in body.split("\n\n"):
    add_para(para)
    add_para("")

add_para("Sincerely:", bold=False)
add_para(f"{recommender_name_en} ({recommender_name_cn})", bold=False)
add_para(f"{recommender_title}, {recommender_affiliation}", bold=False)
contact_line = " | ".join([s for s in [recommender_email, recommender_phone] if s])
if contact_line:
    add_para(contact_line, bold=False)
add_para("")
add_para(f"Applicant: {name_cn} / {name_en} | GPA: {gpa} | Major: {major}", bold=False)
add_para(f"Email: {email} | Phone: {phone}", bold=False)

out_docx = os.path.join(out_dir, f"Recommendation_{name_en.replace(' ', '_')}.docx")
doc.save(out_docx)

pdf_path = out_docx.replace(".docx", ".pdf")
try:
    from docx2pdf import convert
    convert(out_docx, pdf_path)
    print("Generated:", out_docx)
    print("Generated:", pdf_path)
except Exception as e:
    print("DOCX generated, but PDF conversion failed:", e)
