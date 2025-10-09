from docx import Document
from pathlib import Path

def create_template(path: str):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading('Statement of Purpose', level=1)
    doc.add_paragraph('Date: {{ today }}')
    doc.add_paragraph('Applicant: {{ applicant_name }}')
    doc.add_paragraph('Email: {{ email }} | Phone: {{ phone }}')
    doc.add_paragraph('Current University: {{ current_university }} | Degree: {{ degree }} | {{ gpa }}')
    doc.add_paragraph('Target University: {{ university_name }}')
    doc.add_paragraph('Research Area: {{ research_area }}')
    doc.add_paragraph('Top Journals in this Area: {{ top_journals_str }}')
    doc.add_paragraph('Relevant Skills: {{ skills_str }}')
    doc.add_paragraph(' ')
    doc.add_paragraph('Background:')
    doc.add_paragraph('{{ background_summary }}')
    doc.add_paragraph(' ')
    doc.add_paragraph('Motivation:')
    doc.add_paragraph('I am applying to {{ university_name }} to pursue graduate studies focusing on {{ research_area }}. I am particularly inspired by the department\'s strengths and its publication record in {{ top_journals_str }}. My skills in {{ skills_str }} equip me to contribute meaningfully to your research community.')
    doc.add_paragraph(' ')
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('{{ applicant_name }}')
    doc.save(path)

if __name__ == '__main__':
    create_template(str(Path(__file__).resolve().parents[1] / 'template' / 'application_template.docx'))
